from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

def defineBorda(tabela):
    # Define um estilo de borda p/ a tabela
    tbl = tabela._tbl
    for cell in tbl.iter(qn('w:tc')):
        tcPr = cell.get_or_add_tcPr()
        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                              r'<w:top w:val="single" w:sz="12"/>'
                              r'<w:left w:val="single" w:sz="12"/>'
                              r'<w:bottom w:val="single" w:sz="12"/>'
                              r'<w:right w:val="single" w:sz="12"/>'
                              r'</w:tcBorders>')
        tcPr.append(tcBorders)
def centralizaTotal(tabela_destino):
    # Centraliza todas as células da Tabela
    # Itera sobre cada linha da tabela
    for i, row in enumerate(tabela_destino.rows):
        for cell in row.cells:
            # Centraliza horizontalmente o conteúdo da célula
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Deixa em negrito apenas a primeira linha
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
def centralizaPrimeiraLinha(tabela_destino):
    # Centraliza o conteúdo das células da primeira linha (linha 0) horizontalmente
    for cell in tabela_destino.rows[0].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True