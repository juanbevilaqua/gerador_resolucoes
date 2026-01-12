from docx import Document
from docx.shared import Pt

def add_texto_negrito(paragraph, text, bold_word):
    # divide em antes, a palavra(trecho) e depois
    before, middle, after = text.partition(bold_word)

    paragraph.add_run(before)  # parte normal
    paragraph.add_run(middle).bold = True  # parte em negrito
    paragraph.add_run(after)   # parte normal

def add_texto_sublinhado(paragraph, text, bold_word):
    before, middle, after = text.partition(bold_word)

    paragraph.add_run(before)  # parte normal
    paragraph.add_run(middle).underline = True
    paragraph.add_run(after)

def add_texto_italico(paragraph, text, bold_word):
    before, middle, after = text.partition(bold_word)

    paragraph.add_run(before)  # parte normal
    paragraph.add_run(middle).italic = True
    paragraph.add_run(after)

def add_lista_nao_ordenada(doc, itens):
    for item in itens:
        p = doc.add_paragraph()
        p.add_run("     • ").bold = True
        p.add_run(item)

def add_lista_ordenada(doc, itens):
    nums = {
        '1': 'I',
        '2': 'II',
        '3': 'III',
        '4': 'IV',
        '5': 'V',
        '6': 'VI',
        '7': 'VII',
        '8': 'VIII',
        '9': 'IX',
        '10': 'X'
    }

    for i, item in enumerate(itens):
        p = doc.add_paragraph()
        p.add_run(f"     {nums[str(i+1)]}. ").bold = True
        p.add_run(item)
        p_format = p.paragraph_format
        p_format.space_after = Pt(10)
