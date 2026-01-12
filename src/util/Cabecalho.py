from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import yaml

def geraCabecalho(document, ad_referendum, data_reuniao):
    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        configs = list(yaml.safe_load_all(file))

    #vice_coordenador = configs[1]['vice_coordenador']
    funcao = 'EM EXERCÍCIO' if configs[1]['vice_coordenador'] else ''

    modalidade_reunião = "ordinária" if configs[1]["extraordinaria"] is False else "extraordinária"

    if ad_referendum:
        cabecalho = document.add_paragraph(f'        O(A) COORDENADOR(A) {funcao} DO PROGRAMA DE PÓS-GRADUAÇÃO STRICTO SENSU EM CIÊNCIA E TECNOLOGIA AMBIENTAL, da Fundação Universidade '
                                           'Federal da Grande Dourados, no uso de suas atribuições legais, resolve ')
        run = cabecalho.add_run('ad referendum: ')
        run.italic = True
        run.bold = True

    else:
        cabecalho = document.add_paragraph(f'        A COORDENADORIA DO PROGRAMA DE PÓS-GRADUAÇÃO STRICTO SENSU EM CIÊNCIA E TECNOLOGIA AMBIENTAL, da Fundação Universidade Federal da Grande Dourados, '
                                   f'no uso de suas atribuições legais, em reunião {modalidade_reunião} realizada em {data_reuniao}, resolve:')

    cabecalho.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cabecalho_format = cabecalho.paragraph_format
    cabecalho_format.space_after = Pt(10)
