from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTabela, CarregadorDeConfigs
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, valores_dinamicos):
    nivel_discente = valores_dinamicos["Nível do Discente"]
    nome = valores_dinamicos["Nome do Discente"]
    projeto_atual = valores_dinamicos["Projeto Atual"]
    novo_projeto = valores_dinamicos["Novo Projeto"]

    file_parts = CarregadorDeConfigs.carregar_config()
    document = Document(str(file_parts[0]['timbre_res']))

    #n_res, data_res, ad_referendum, data_reuniao, nivel_discente, nome, projeto_atual, novo_projeto = ColetorDeDados.coletaDados(15)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum,data_reuniao)


    p1 = document.add_paragraph(f'   APROVAR a troca de projeto de pesquisa do(a) discente de {nivel_discente}, ')

    p1.add_run(f'{nome}').bold = True
    p1.add_run(', conforme segue: ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(20)

    # Gera tabela
    tabela = document.add_table(rows=2, cols=2)

    # Definir um estilo de borda básica para cada célula
    FormatadorTabela.defineBorda(tabela)

    tabela.cell(0, 0).text = 'Projeto Atual:'
    tabela.cell(0, 1).text = projeto_atual

    tabela.cell(1, 0).text = 'Novo Projeto: '
    tabela.cell(1, 1).text = novo_projeto

    # Centraliza o conteúdo das células horizontalmente
    for row in tabela.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova troca de projeto de pesquisa - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova troca de projeto de pesquisa - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)