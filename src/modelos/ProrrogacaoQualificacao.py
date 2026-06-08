from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, CarregadorDeConfigs
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    nivel_discente = dados_dinamicos["Nível do Discente"]
    nome = dados_dinamicos["Nome do Discente"]
    ano_ingresso = dados_dinamicos["Ano de Ingresso"]
    data_limite = dados_dinamicos["Data Limite Aprovada"]

    file_parts = CarregadorDeConfigs.carregar_config()
    document = Document(str(file_parts[0]['timbre_res']))

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'      APROVAR o novo prazo de qualificação para o(a) discente de {nivel_discente}, ')

    p1.add_run(f'{nome}').bold = True
    p1.add_run(f'(turma de ingresso PPGCTA-{ano_ingresso}), tendo como data limite o dia ')
    p1.add_run(f'{data_limite}.').bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(100)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova prorrogação de qualificação - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova prorrogação de qualificação - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)