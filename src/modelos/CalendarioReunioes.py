from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTabela
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos, configs):
    ano = dados_dinamicos["Ano Vigente"]
    reunioes = dados_dinamicos["Nº da Reunião"]
    cont_reunioes = len(reunioes)
    data_reunioes = dados_dinamicos["Data da Reunião"]

    #document = Document('MODELO papel timbrado FACET.docx')
    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))

    #n_res, data_res, ad_referendum, data_reuniao, ano, cont_reunioes, reunioes, data_reunioes  = ColetorDeDados.coletaDados(8)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'       Aprovar o calendário de reuniões ordinárias de {ano}, do PPGCTA, conforme segue: ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(10)

    tabela = document.add_table(rows=cont_reunioes+1, cols=2)# rows=cont_alunos+1, pois cada aluno precisa de uma linha + a linha de identif. dos campos(discentes, disciplinas,etc)

    # Definir um estilo de borda básica para cada célula
    FormatadorTabela.defineBorda(tabela)

    tabela.cell(0, 0).text = 'REUNIÃO'
    tabela.cell(0, 1).text = 'DATA PREVISTA'

    for i in range(cont_reunioes):# gera uma nova linha de tabela para cada reunião e referencia suas informações
        tabela.cell(i+1, 0).text = reunioes[i]# incrementa 1 na linha pois a linha 0 já possui os identificadores dos campos
        tabela.cell(i+1, 1).text = data_reunioes[i]

    FormatadorTabela.centralizaTotal(tabela)

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list): # se for uma lista, republicacao == True
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = ColetorDeDados.extraiAnoResolucao(data_res)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova calendário de reuniões ordinárias {ano}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova calendário de reuniões ordinárias {ano}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)