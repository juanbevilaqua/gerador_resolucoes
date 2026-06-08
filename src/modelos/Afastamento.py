from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, CarregadorDeConfigs
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    nivel_discente = dados_dinamicos["Nível do Discente"]
    nome = dados_dinamicos["Nome do Discente"]
    n_dias_afast = dados_dinamicos["Nº. de Dias de Afastamento"]
    data_inicio = dados_dinamicos["Data de Início"]
    data_fim = dados_dinamicos["Data de Finalização"]
    motivo = dados_dinamicos["Motivo"]
    if motivo == "Particular":
        motivo = "motivos particulares do(a) discente."
    elif motivo == "Saúde":
        motivo = "questões de saúde, comprovadas através de atestado médico apresentado."
    else:
        outro = dados_dinamicos["Outro"]
        motivo = outro


    file_parts = CarregadorDeConfigs.carregar_config()
    document = Document(str(file_parts[0]['timbre_res']))

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'      I. Aprovar o afastamento do(a) discente de {nivel_discente}, ')
    p1.add_run(f'{nome}').bold = True
    p1.add_run(f', por {n_dias_afast} dias, justificado por {motivo}.')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p2 = document.add_paragraph(f'      II. Informar a suspensão das atividades do(a) discente no programa no período de ')
    p2.add_run(f'{data_inicio} a {data_fim}.').bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(100)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):  # se for uma lista, republicacao == True
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova afastamento - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova afastamento - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)

