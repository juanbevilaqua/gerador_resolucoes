from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTexto
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    nome = dados_dinamicos["Nome do Professor"]
    data_inicio_relatorio = dados_dinamicos["Data de Início das Atividades"]
    data_fim_relatorio = dados_dinamicos["Data de Finalização das Atividades"]
    tipo_relatorio = dados_dinamicos["Tipo de Aprovação"]

    if tipo_relatorio == "Com Renovação":
        data_inicio_renovacao = dados_dinamicos["Data de Início da Renovação"]
        data_fim_renovacao = dados_dinamicos["Data Final da Renovação"]
        supervisor = dados_dinamicos["Professor Supervisor"]

    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))
    #n_res, data_res, ad_referendum, data_reuniao, nivel_discente, nome, ano_ingresso, data_fim_lic, data_ini_defesa, data_lmt_defesa = ColetorDeDados.coletaDados(12)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    titulo_parcial = "Aprova relatorio de atividades(Professor Visitante) - "

    p1 = document.add_paragraph()
    FormatadorTexto.add_texto_sublinhado(p1, f"     I - APROVAR o relatório de atividades desenvolvidas no PPGCTA, no período de {data_inicio_relatorio} a {data_fim_relatorio},",
                                         f"{data_inicio_relatorio} a {data_fim_relatorio}")
    FormatadorTexto.add_texto_negrito(p1, f" referente ao(à) professor(a) visitante: {nome}.", f"{nome}")
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(15)

    if tipo_relatorio == "Com Renovação":

        titulo_parcial = "Aprova relatorio de atividades e renovação de contrato(Professor Visitante) - "

        p2 = document.add_paragraph()
        FormatadorTexto.add_texto_sublinhado(p2, f"     II - MANIFESTAR-SE FAVORAVELMENTE à renovação de contrato do(a) professor(a) visitante, com vigência de {data_inicio_renovacao} a {data_fim_renovacao}.", f"{data_inicio_renovacao} a {data_fim_renovacao}")
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2_format = p2.paragraph_format
        p2_format.space_after = Pt(15)

        p3 = document.add_paragraph()
        FormatadorTexto.add_texto_sublinhado(p3, f"      III - Aprovar o novo plano de trabalho do(a) docente, a ser desempenhado no período de {data_inicio_renovacao} a {data_fim_renovacao}, sob supervisão de {supervisor}.", f"{data_inicio_renovacao} a {data_fim_renovacao}")
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p4 = document.add_paragraph()
    p4_format = p4.paragraph_format
    p4_format.space_after = Pt(100)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = ColetorDeDados.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM {titulo_parcial} - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - {titulo_parcial} - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)

