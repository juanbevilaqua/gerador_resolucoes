from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTabela
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):

    lingua = dados_dinamicos["Língua"]
    discentes = dados_dinamicos["Nome do Discente"]
    cont_discentes = len(discentes)
    datas = dados_dinamicos["Data do Exame"]
    resolucoes = dados_dinamicos["Resolução de Aprovação"]


    # self.criar_campo("Data do Exame")
    # self.criar_campo("Resolução de Aprovação")
    # lingua, discentes, cont_discentes, datas, resolucoes
    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))

    #n_res, data_res, ad_referendum, data_reuniao, lingua, discentes, cont_discentes, datas, resolucoes = ColetorDeDados.coletaDados(13)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'       Aprovar o aproveitamento do resultado do exame de suficiência em ')
    p1.add_run(f'{lingua}').bold = True
    p1.add_run(', realizado no processo seletivo de mestrado, dos doutorandos a seguir relacionados: ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(10)

    tabela = document.add_table(rows=cont_discentes+1, cols=4)# rows=cont_alunos+1, pois cada aluno precisa de uma linha + a linha de identif. dos campos(discentes, disciplinas,etc)

    # Definir um estilo de borda básica para cada célula
    FormatadorTabela.defineBorda(tabela)

    tabela.cell(0, 0).text = 'DISCENTE'
    tabela.cell(0, 1).text = 'DATA DO EXAME'
    tabela.cell(0, 2).text = 'RESULTADO'
    tabela.cell(0, 3).text = 'RESOLUÇÃO'

    for i in range(cont_discentes):# gera uma nova linha de tabela para cada aluno e referencia suas informações
        tabela.cell(i+1, 0).text = discentes[i]# incrementa 1 na linha pois a linha 0 já possui os identificadores dos campos
        tabela.cell(i+1, 1).text = datas[i]
        tabela.cell(i+1, 2).text = 'APROVADO(A)'
        tabela.cell(i+1, 3).text = resolucoes[i]

    FormatadorTabela.centralizaTotal(tabela)

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):  # se for uma lista, republicacao == True
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    #conj_nomes_encurtados = ', '.join(ColetorDeDados.encurtaNome(discente) for discente in discentes)

    if ad_referendum:
        #titulo_doc = f'Resolução nº {n_res}  - AD REFERENDUM Aprova aproveitamento de suficiência em {lingua} - {conj_nomes_encurtados}.docx'
        titulo_doc = f'Resolução nº {n_res}  - AD REFERENDUM Aprova aproveitamento de suficiência em {lingua} - {discentes}.docx'

    else:
        #titulo_doc = f'Resolução nº {n_res}  - Aprova aproveitamento de suficiência em {lingua} - {conj_nomes_encurtados}.docx'
        titulo_doc = f'Resolução nº {n_res}  - Aprova aproveitamento de suficiência em {lingua} - {discentes}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)