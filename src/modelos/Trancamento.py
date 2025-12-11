from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos, configs):
    nivel_discente = dados_dinamicos["Nível do Discente"]
    nome = dados_dinamicos["Nome do Discente"]
    rga = dados_dinamicos["RGA"]
    semestre = dados_dinamicos["Semestre de Trancamento(ano-semestre)"]
    motivo = dados_dinamicos["Motivo"]

    if motivo == "Outro(s)":
        motivo = dados_dinamicos["Outro(s)"]


    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))
    #n_res, data_res, ad_referendum, data_reuniao, nivel_discente, nome, rga, semestre, motivo = ColetorDeDados.coletaDados(2)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph('      APROVAR o trancamento do semestre ')
    p1.add_run(f'{semestre}').bold = True
    p1.add_run(f', a pedido, do(a) acadêmico(a)')
    p1.add_run(f'{nome}').bold = True
    p1.add_run(f'(RGA: {rga}), por {motivo}.')
    #p1.add_run(f'{data_limite}.').bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(100)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    dir_res = ColetorDeDados.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova trancamento do semestre - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova trancamento do semestre - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)