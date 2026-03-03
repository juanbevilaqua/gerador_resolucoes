from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTexto
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    nome = dados_dinamicos["Nome do Professor"]
    classificacao = dados_dinamicos["Classificação no Processo Seletivo"]
    edital_propp = dados_dinamicos["Nº Edital PROPP"]
    data_edital_propp = dados_dinamicos["Data Publicação Edital PROPP"]
    cadastro_reserva = dados_dinamicos["Cadastro de Reserva?"]

    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))
    #n_res, data_res, ad_referendum, data_reuniao, nivel_discente, nome, exame, lingua = ColetorDeDados.coletaDados(14)# 1 indica o tipo de resolução

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'      APROVAR a contratação do(a) docente {nome}, classificada em {classificacao}º lugar no Processo Seletivo Simplificado para '
                                f'Contratação de Professor(a) Visitante Júnior ')

    if cadastro_reserva == "Sim":
        p1.add_run('(cadastro de reserva)')

    p1.add_run(f', do PPGCTA (edital de homologação PROPP nº {edital_propp}, de {data_edital_propp}).')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(100)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova contratação de professor visitante - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova contratação de professor visitante - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)

