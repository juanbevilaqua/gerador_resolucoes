from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTabela, FormatadorTexto
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    processo_seletivo = dados_dinamicos["Processo Seletivo"]
    outro = dados_dinamicos["Outro"]
    if outro != 'null':
        processo_seletivo = outro

    semestre = dados_dinamicos["Semestre(ano-nº semestre)"]
    vagas_mtd = dados_dinamicos["Vagas de Mestrado"]
    vagas_dtd = dados_dinamicos["Vagas de Doutorado"]

    cont_col = 0

    if vagas_mtd != '0':
        cont_col += 1

    if vagas_dtd != '0':
        cont_col += 1

    parts_semestre = semestre.split(".") # Formato: 2025.2
    if parts_semestre[-1] == "1":
        semestre_convert = f"primeiro semestre de {parts_semestre[0]}"
    else:
        semestre_convert = f"segundo semestre de {parts_semestre[0]}"

    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph()
    FormatadorTexto.add_texto_sublinhado(p1, f"     APROVAR, o número de vagas para o {processo_seletivo} do {semestre_convert}, conforme segue:", f"{semestre_convert}")
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(10)

    # itens que serão inseridos na lista(não ordenada, nesse caso)
    list_itens = []

    if vagas_mtd != '0' and vagas_dtd == '0':
        list_itens.append(f'MESTRADO: {vagas_mtd} vagas')

    elif vagas_dtd != '0' and vagas_mtd == '0':
        list_itens.append(f'DOUTORADO: {vagas_dtd} vagas')

    elif vagas_dtd != '0' and vagas_mtd != '0':
        list_itens.append(f'MESTRADO: {vagas_mtd} vagas')
        list_itens.append(f'DOUTORADO: {vagas_dtd} vagas')


    FormatadorTexto.add_lista_nao_ordenada(document, list_itens)

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova número de vagas  {semestre}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova número de vagas {processo_seletivo} - {semestre}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)