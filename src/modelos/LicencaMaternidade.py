from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos, configs):
    nivel_discente = dados_dinamicos["Nível do Discente"]
    nome = dados_dinamicos["Nome do Discente"]
    qtde_meses = dados_dinamicos["Nº. de Meses da Licença(por extenso. Ex: seis)"]
    ano_ingresso = dados_dinamicos["Ano de Ingresso"]
    data_fim_lic = dados_dinamicos["Data Fim da Licença"]
    data_ini_defesa = dados_dinamicos["Data Inicial de Defesa"]
    data_lmt_defesa = dados_dinamicos["Data Limite Para Defesa Ajustada"]

    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))
    #n_res, data_res, ad_referendum, data_reuniao, nivel_discente, nome, ano_ingresso, data_fim_lic, data_ini_defesa, data_lmt_defesa = ColetorDeDados.coletaDados(12)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'      1. APROVAR, por {qtde_meses} meses, a contar da data de nascimento do criança, a licença maternidade para a discente de {nivel_discente}, ')
    p1.add_run(f'{nome}').bold = True
    p1.add_run(f'(turma de ingresso PPGCTA/{ano_ingresso}), até a data de {data_fim_lic};')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p2 = document.add_paragraph(f'      2. INFORMAR que este período não será contabilizada no tempo de curso;')

    p3 = document.add_paragraph(f'      3. ADEQUAR a data de defesa, antes prevista para  ')
    p3.add_run(f'{data_ini_defesa}').underline = True
    p3.add_run(f', para a data limite de ')
    p3.add_run(f'{data_lmt_defesa}').bold = True
    p3.add_run(f', com base no período da licença.')
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p3_format = p3.paragraph_format
    p3_format.space_after = Pt(100)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = ColetorDeDados.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova licença maternidade e ajusta data(s) - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova licença maternidade e ajusta data(s) - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)

