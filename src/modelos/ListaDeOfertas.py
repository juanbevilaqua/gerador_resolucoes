from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTabela, FormatadorTexto
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml
from src.controladores.controladorDisciplina import DisciplinaController

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    ano_semestre = dados_dinamicos["ano-semestre"]
    nomes_disciplinas = dados_dinamicos["Disciplina"]
    docentes = dados_dinamicos["Professor Responsável"]
    print("Docentes lisa", docentes)
    conj_docentes_total = []
    for grupo in docentes:
        conj_docentes_parcial = ', '.join(grupo)
        conj_docentes_total.append(conj_docentes_parcial)
    #conj_docentes = ', '.join(docente for docente in docentes)
    cont_disciplinas = len(nomes_disciplinas)
    #data_reunioes = dados_dinamicos["Data da Reunião"]

    dados_disciplinas = []
    print("Docentes lisa", docentes)

    for nome in nomes_disciplinas:
        disciplina = DisciplinaController.buscar_por_nome(nome)
        dados_disciplinas.append(disciplina[0])

    print("Dados disciplina: ", dados_disciplinas)

    #document = Document('MODELO papel timbrado FACET.docx')
    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))

    #n_res, data_res, ad_referendum, data_reuniao, ano, cont_reunioes, reunioes, data_reunioes  = ColetorDeDados.coletaDados(8)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph()#(f'       Aprovar a Lista de Ofertas de disciplinas do PPGCTA para {ano_semestre}, conforme segue: ')
    FormatadorTexto.add_texto_negrito(p1, f"     APROVAR a Lista de Ofertas de disciplinas do PPGCTA para {ano_semestre}, conforme segue: ", ano_semestre )

    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(10)

    tabela = document.add_table(rows=cont_disciplinas+1, cols=4)# rows=cont_alunos+1, pois cada aluno precisa de uma linha + a linha de identif. dos campos(discentes, disciplinas,etc)

    # Definir um estilo de borda básica para cada célula
    FormatadorTabela.defineBorda(tabela)

    tabela.autofit = False
    tabela.columns[0].width = Inches(3)
    tabela.columns[1].width = Inches(2)
    tabela.columns[2].width = Inches(1)
    tabela.columns[3].width = Inches(1.5)

    tabela.cell(0, 0).text = 'DISCIPLINA'
    tabela.cell(0, 1).text = 'DOCENTE'
    tabela.cell(0, 2).text = 'C.H.(h/a)'
    tabela.cell(0, 3).text = 'CRÉDITOS'

    for i in range(cont_disciplinas):# gera uma nova linha de tabela para cada reunião e referencia suas informações
        tabela.cell(i+1, 0).text = dados_disciplinas[i][1]# NOME DISCIPLINA
        tabela.cell(i+1, 1).text = conj_docentes_total[i] # DOCENTE
        tabela.cell(i+1, 2).text = str(dados_disciplinas[i][2]) # C.H.
        tabela.cell(i+1, 3).text = str(dados_disciplinas[i][3])  # CRÉDITOS

    FormatadorTabela.centralizaTotal(tabela)

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list): # se for uma lista, republicacao == True
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = ColetorDeDados.extraiAnoResolucao(data_res)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova lista de ofertas {ano_semestre}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova lista de ofertas {ano_semestre}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)