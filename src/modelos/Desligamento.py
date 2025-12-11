from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTabela
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos, configs):
    niveis_discentes = dados_dinamicos["Nível do Discente"]
    discentes = dados_dinamicos["Nome do Discente"]
    cont_discentes = len(discentes)

    motivos = dados_dinamicos["Motivo"]
    for i, motivo in enumerate(motivos):
        if motivo == "Não realização de matrícula":
            ano_semestre = dados_dinamicos["Complemento"][i]
            motivos[i] = f"Matrícula não realizada no semestre {ano_semestre}(item IV, artigo 45 do regulamento do PPGCTA de 2023)"
        elif motivo == "Outro(s)":
            outro = dados_dinamicos["Complemento"][i]
            motivos[i] = outro

    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))

    #n_res, data_res, ad_referendum, data_reuniao, discentes, cont_discentes, niveis_discentes, motivos = ColetorDeDados.coletaDados(3)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph('      APROVAR o desligamento dos discentes do PPGCTA, conforme segue: ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(15)

    tabela = document.add_table(rows=cont_discentes+1, cols=3)# rows=cont_discentes+1, pois cada aluno precisa de uma linha + a linha de identif. dos campos(discentes, disciplinas,etc)

    # Definir um estilo de borda básica para cada célula
    FormatadorTabela.defineBorda(tabela)

    # # Definir a largura das colunas da tabela (em twips, 1 polegada = 1440 twips)
    # tabela.columns[0].width = 200000  # Largura da primeira coluna (exemplo: DISCENTE)
    # tabela.columns[1].width = 375000  # Largura da segunda coluna (exemplo: DISCIPLINA(S))
    # tabela.columns[2].width = 275000  # Largura da terceira coluna (exemplo: DOCENTE RESPONSÁVEL)
    # #tabela.columns[3].width = 50000  # Largura da quarta coluna (exemplo: SEMESTRE DE OFERTA)

    tabela.cell(0, 0).text = 'DISCENTE'
    tabela.cell(0, 1).text = 'NÍVEL'
    tabela.cell(0, 2).text = 'MOTIVO'

    for i in range(cont_discentes):# gera uma nova linha de tabela para cada aluno e referencia suas informações
        tabela.cell(i+1, 0).text = discentes[i]# incrementa 1 na linha pois a linha 0 já possui os identificadores dos campos
        tabela.cell(i+1, 1).text = niveis_discentes[i]
        tabela.cell(i+1, 2).text = motivos[i]

    FormatadorTabela.centralizaPrimeiraLinha(tabela)

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = ColetorDeDados.extraiAnoResolucao(data_res)
    conj_nomes_encurtados = ', '.join(ColetorDeDados.encurtaNome(discente) for discente in discentes)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res}  - AD REFERENDUM Aprova desligamento(s) - {conj_nomes_encurtados}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res}  - Aprova desligamento(s) - {conj_nomes_encurtados}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)