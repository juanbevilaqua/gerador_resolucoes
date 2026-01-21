from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTabela
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    semestre = dados_dinamicos["ano-semestre"]
    niveis_discentes = dados_dinamicos["Nível do Discente"]
    discentes = dados_dinamicos["Nome do Discente"]
    cont_discentes = len(discentes)

    conj_geral_disciplinas = dados_dinamicos["Disciplina"]#todas as disciplinas de todos os alunos
    conj_geral_professores = dados_dinamicos["Professor"]

    # ajuste do formato de exibição dos elementos(disciplinas e professores) na tabela
    conj_geral_disciplinas_ajust = []
    conj_geral_professores_ajust = []

    for disciplinas in conj_geral_disciplinas:
       disciplinas_ajust = ', '.join(f"{j+1}. {disciplina}" for j, disciplina in enumerate(disciplinas))# Ex: 1. Disciplina X, 2. Disciplina Y
       conj_geral_disciplinas_ajust.append(disciplinas_ajust)

    for professores in conj_geral_professores:
       professores_ajust = ', '.join(f"{j+1}. {professor}" for j, professor in enumerate(professores))
       conj_geral_professores_ajust.append(professores_ajust)

    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))

    #n_res, data_res, ad_referendum, data_reuniao, nivel_discente, discentes, todas_disciplinas, todos_docentes, cont_discentes, niveis_discentes = ColetorDeDados.coletaDados(4)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'      APROVAR os pedidos de cancelamento de matrícula nas disciplinas no semestre {semestre}, conforme segue: ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(20)

    tabela = document.add_table(rows=cont_discentes+1, cols=3)# rows=cont_discentes+1, pois cada aluno precisa de uma linha + a linha de identif. dos campos(discentes, disciplinas,etc)

    # Definir um estilo de borda básica para cada célula
    FormatadorTabela.defineBorda(tabela)

    tabela.autofit = False#desabitita autoajuste do word para evitar desformatação

    # # Definir a largura das colunas da tabela
    tabela.columns[0].width = Inches(1.5)  # DISCENTE
    tabela.columns[1].width = Inches(3)  # DISCIPLINA(S)
    tabela.columns[2].width = Inches(2)

    # tabela.columns[0].width = 200000  # Largura da primeira coluna (DISCENTE)
    # tabela.columns[1].width = 375000  # Largura da segunda coluna (DISCIPLINA(S))
    # tabela.columns[2].width = 275000  # Largura da terceira coluna (DOCENTE RESPONSÁVEL)
    # #tabela.columns[3].width = 50000  # Largura da quarta coluna (SEMESTRE DE OFERTA)

    tabela.cell(0, 0).text = 'DISCENTE'
    tabela.cell(0, 1).text = 'DISCIPLINA(S)'
    tabela.cell(0, 2).text = 'DOCENTE(S) RESPONSÁVEL(EIS)'
    #tabela.cell(0, 3).text = 'SEMESTRE DE OFERTA'

    for i in range(cont_discentes):# gera uma nova linha de tabela para cada aluno e referencia suas informações
        tabela.cell(i+1, 0).text = discentes[i] + ' (' + niveis_discentes[i] + ')'# incrementa 1 na linha pois a linha 0 já possui os identificadores dos campos
        tabela.cell(i+1, 1).text = conj_geral_disciplinas_ajust[i]
        tabela.cell(i+1, 2).text = conj_geral_professores_ajust[i]
        #tabela.cell(i+1, 3).text = '2025/2'

    FormatadorTabela.centralizaPrimeiraLinha(tabela)

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = ColetorDeDados.extraiAnoResolucao(data_res)
    conj_nomes_encurtados = ', '.join(ColetorDeDados.encurtaNome(discente) for discente in discentes)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res}  - AD REFERENDUM Aprova cancelamento de matrícula em disciplinas - {conj_nomes_encurtados}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res}  - Aprova cancelamento de matrícula em disciplinas - {conj_nomes_encurtados}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)