from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ManipuladorDeArquivos, CarregadorDeConfigs, Assinatura, BuscadorDeArquivos, FormatadorTexto
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):

    lista_res = dados_dinamicos["Resolução (nº-Ano)"]
    conj_res = ', '.join(res for res in lista_res)

    file_parts = CarregadorDeConfigs.carregar_config()
    document = Document(str(file_parts[0]['timbre_res']))

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, False, data_reuniao)

    p1 = document.add_paragraph()

    if len(lista_res) > 1:
        FormatadorTexto.add_texto_negrito(p1, f"     Homologar em bloco as resoluções {conj_res}", conj_res)
    else:
        FormatadorTexto.add_texto_negrito(p1, f"     Homologar a resolução nº. {conj_res}", conj_res)

    #p1 = document.add_paragraph(f'      Homologar em bloco as resoluções ')

    #p1.add_run(f'{conj_res}').bold = True
    #p1.add_run(f', emitida(s) Ad Referendum pela coordenação do PPGCTA.')#Todas os documentos citados nesta resolução encontram-se em anexo.'
    FormatadorTexto.add_texto_italico(p1, f', emitida(s) Ad Referendum pela coordenação do PPGCTA.', 'Ad Referendum')
    #p1.add_run(f'{data_limite}.').bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(100)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    #dir_res = ColetorDeDados.extraiAnoResolucao(data_res)

    #titulo_doc = f'Resolução nº {n_res} - Homologa resoluções Ad Referendum - {conj_res}.docx'

    dir_res = util.Data.extraiAnoResolucao(data_res)

    titulo_doc = f'Resolução nº {n_res} - Homologa resoluções Ad Referendum - {conj_res}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc) # a compilação de arquivos necessita da versao pdf

    # ** ESTE TRECHO REALIZA O MERGE DA RES. DE HOMOLOG. COM AS AD REFERENDUNS. FUNCIONAL, MAS EM STAND-BY NO MOMENTO **
    #
    # # Define resoluçoes pra unir
    # lista_caminhos_res = []
    # titulo_pdf = f'./resolucoes/{dir_res}/Resolução nº {n_res} - Homologa resoluções Ad Referendum - {conj_res}.pdf'
    # # Faz o append da resolução de homologação como a primeira do documento unificado
    # lista_caminhos_res.append(titulo_pdf)
    #
    # # Loop para informa das resoluções que serão unidas
    # for res in lista_res:
    #     #titulo_res = input('Informe o título da resolução ad referendum: ')
    #     # partes = res.split('-')
    #     # caminho_res = f'./resolucoes/{partes[-1]}/{titulo_res}'# partes[-1] representa o ano da res.Ex: 100-2022 --> 2022
    #     caminho_res = BuscadorDeArquivos.buscarResolucao(res)
    #     lista_caminhos_res.append(caminho_res)
    #
    # ManipuladorDeArquivos.unirPdfs(lista_caminhos_res, titulo_pdf)# sobscreva a resolução individual p/ o documento unificado