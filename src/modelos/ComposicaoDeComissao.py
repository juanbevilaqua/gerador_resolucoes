from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, FormatadorTabela
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    comissao = dados_dinamicos["Nome da Comissão"]
    membros = dados_dinamicos["Professor Membro"]
    cont_membros = len(membros)
    tipos_part = dados_dinamicos["Tipo de Participação"]

    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))

    #n_res, data_res, ad_referendum, data_reuniao, comissao, cont_membros, membros, tipos_part = ColetorDeDados.coletaDados(16)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'       APROVAR a composição da ')
    p1.add_run(f'{comissao}').bold = True
    p1.add_run(', do PPGCTA, visando a condução das etapas do certame. Composição definida com os seguintes membros:')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(10)

    tabela = document.add_table(rows=cont_membros+1, cols=2)

    FormatadorTabela.defineBorda(tabela)

    tabela.cell(0, 0).text = 'MEMBRO'
    tabela.cell(0, 1).text = 'TIPO DE PARTICIPAÇÃO'

    for i in range(cont_membros):# gera uma nova linha de tabela para cada membro e referencia suas informações
        tabela.cell(i+1, 0).text = membros[i]# incrementa 1 na linha pois a linha 0 já possui os identificadores dos campos
        tabela.cell(i+1, 1).text = tipos_part[i]

    FormatadorTabela.centralizaTotal(tabela)

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = ColetorDeDados.extraiAnoResolucao(data_res)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova composição da {comissao}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova composição da {comissao}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)