from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ManipuladorDeArquivos, ColetorDeDados, Assinatura, FormatadorTabela
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):


    nivel_discente = dados_dinamicos["Nível do Discente"]
    nome = dados_dinamicos["Nome do Discente"]
    tipo_trabalho = dados_dinamicos["Tipo de Trabalho"]
    titulo_trabalho = dados_dinamicos["Título do Trabalho"]
    tipo_apresentacao = dados_dinamicos["Tipo de Apresentação"]
    data_apresentacao = dados_dinamicos["Data da Apresentação"]

    #doc_pdf = input('Informe a solicitação de banca: ')7
    # Carrega o documento de origem
    #doc_origem = ManipuladorDeArquivos.converterPdfDocx('solicitacao_banca.pdf')

    doc_origem = ManipuladorDeArquivos.converterPdfDocx(dados_dinamicos["Solicitação de Banca"])
    # Carrega ou cria o documento de destino
    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))

    #Seleciona a tabela 1 do documento padrão
    tabela_origem = doc_origem.tables[1]#extraindo a tabela 1 do arquivo gerado pelo SCPG
    #Caso exista um erro de interpretação, pega a tabela seguinte(é possível que o título do trabalho seja identificado como uma tabela)
    if len(tabela_origem.rows) <= 2:
        tabela_origem = doc_origem.tables[2]
    print("NÚMERO DE LINHAS TABELA***:", len(tabela_origem.rows))

    #n_res, data_res, ad_referendum, data_reuniao, nivel_discente, nome, tipo_trabalho, titulo_trabalho, tipo_apresentacao, data_apresentacao = ColetorDeDados.coletaDados(7)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'      Aprovar a banca de {tipo_apresentacao} da {tipo_trabalho}:')

    p1.add_run(f' “{titulo_trabalho}”').italic = True
    p1.add_run(f', do(a) discente de {nivel_discente}, ')
    p1.add_run(f'{nome}').bold = True
    p1.add_run(', conforme segue: ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(10)

    # Cria uma nova tabela no documento de destino com o mesmo número de linhas e colunas
    tabela_destino = document.add_table(rows=len(tabela_origem.rows), cols=len(tabela_origem.columns))

    # Definir um estilo de borda básica para cada célula
    FormatadorTabela.defineBorda(tabela_destino)


    # Copia cada célula da tabela de origem para a tabela de destino
    for i, linha in enumerate(tabela_origem.rows):
        for j, cell in enumerate(linha.cells):
            tabela_destino.cell(i, j).text = cell.text

    FormatadorTabela.centralizaPrimeiraLinha(tabela_destino)

    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    p2_format.space_after = Pt(5)

    p3 = document.add_paragraph(f'Data da {tipo_apresentacao}: {data_apresentacao}')

    p4 = document.add_paragraph()
    p4_format = p4.paragraph_format
    p4_format.space_after = Pt(80)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):  # se for uma lista, republicacao == True
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova banca de {tipo_apresentacao} - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova banca de {tipo_apresentacao} - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)