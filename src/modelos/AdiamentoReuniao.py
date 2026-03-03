from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Cabecalho import geraCabecalho
from src.util.Titulo import geraTitulo
from src.util.RodapeRepublicacao import geraRodapeRepublicacao
from src.util import Armazenador, ColetorDeDados, Assinatura, Data
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):


    n_reuniao = dados_dinamicos["Nº. da Reunião"]
    data_inicial = dados_dinamicos["Data Inicial"]
    resolucao = dados_dinamicos["Resolução de Aprovação"]
    data_res_original = dados_dinamicos["Data da Resolução Original"]
    nova_data = dados_dinamicos["Data Atualizada"]
    if nova_data == "null":
        previsao = "Sine Die"
    else:
        previsao = f"para {nova_data}"

    with open('./src/config/configs.yaml', "r", encoding="utf-8") as file:
        file_parts = list(yaml.safe_load_all(file))
    document = Document(str(file_parts[0]['timbre_res']))
    #n_res, data_res, ad_referendum, data_reuniao, n_reuniao, data_inicial, resolucao, data_res_original, previsao = ColetorDeDados.coletaDados(9)# 1 indica o tipo de resolução, nesse caso, Prorrogação de Qualificação

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'      Adiar {previsao} a {n_reuniao}ª. reunião ordinária da Coordenadoria do PPGCTA, prevista para realizar-se inicialmente em {data_inicial}')


    p1.add_run(f' (RESOLUÇÃO {resolucao}, de {data_res_original}).').bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(100)


    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):  # se for uma lista, republicacao == True
        geraRodapeRepublicacao(document)

    # Define o diretório e título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Adia {n_reuniao}a reunião ordinária.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Adia {n_reuniao}a reunião ordinária.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)

