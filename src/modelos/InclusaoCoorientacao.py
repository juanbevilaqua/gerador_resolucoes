from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import util.Data
from src.util.Titulo import geraTitulo
from src.util.Cabecalho import geraCabecalho
from src.util import Armazenador, ColetorDeDados, Assinatura, CarregadorDeConfigs
from util.RodapeRepublicacao import geraRodapeRepublicacao
import yaml

def geraModelo(n_res, data_res, ad_referendum, data_reuniao, dados_dinamicos):
    nivel_discente = dados_dinamicos["Nível do Discente"]
    nome = dados_dinamicos["Nome do Discente"]
    coorientador = dados_dinamicos["Coorientador(a)"]
    universidade = dados_dinamicos["Universidade"]

    file_parts = CarregadorDeConfigs.carregar_config()
    document = Document(str(file_parts[0]['timbre_res']))
    #n_res, data_res, ad_referendum, data_reuniao, nivel_discente, nome, coorientador, universidade = ColetorDeDados.coletaDados(11)

    geraTitulo(document, n_res, data_res)

    geraCabecalho(document, ad_referendum, data_reuniao)

    p1 = document.add_paragraph(f'      Aprovar a inclusão do(a) docente ')
    p1.add_run(f'{coorientador}').underline = True
    p1.add_run(f', da {universidade}, como coorientador(a) do(a) discente de {nivel_discente}: ')
    p1.add_run(f'{nome}.').bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(100)

    Assinatura.geraCampoAssinatura(document)

    if isinstance(file_parts[1]['republicacao'], list):
        geraRodapeRepublicacao(document)

    # Define o título da resolução que será salva
    dir_res = util.Data.extraiAnoResolucao(data_res)
    nome_encurtado = ColetorDeDados.encurtaNome(nome)
    if ad_referendum:
        titulo_doc = f'Resolução nº {n_res} - AD REFERENDUM Aprova coorientação - {nome_encurtado}.docx'
    else:
        titulo_doc = f'Resolução nº {n_res} - Aprova coorientação - {nome_encurtado}.docx'

    Armazenador.salvar(dir_res, document, titulo_doc)