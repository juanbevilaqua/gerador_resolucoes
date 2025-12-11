document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    # document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')


#GERAR TABELA COM BORDAS
tabela = document.add_table(rows=3, cols=3)

    # Definir um estilo de borda básica para cada célula

from docx.oxml.ns import qn
from docx.oxml import parse_xml

    tbl = tabela._tbl
    for cell in tbl.iter(qn('w:tc')):
        tcPr = cell.get_or_add_tcPr()
        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                              r'<w:top w:val="single" w:sz="12"/>'
                              r'<w:left w:val="single" w:sz="12"/>'
                              r'<w:bottom w:val="single" w:sz="12"/>'
                              r'<w:right w:val="single" w:sz="12"/>'
                              r'</w:tcBorders>')
        tcPr.append(tcBorders)

    tabela.cell(0, 0).text = 'Cabeçalho 1'
    tabela.cell(0, 1).text = 'Cabeçalho 2'
    tabela.cell(0, 2).text = 'Cabeçalho 3'

    tabela.cell(1, 0).text = 'Linha 1, Coluna 1'
    tabela.cell(1, 1).text = 'Linha 1, Coluna 2'
    tabela.cell(1, 2).text = 'Linha 1, Coluna 3'

    tabela.cell(2, 0).text = 'Linha 2, Coluna 1'
    tabela.cell(2, 1).text = 'Linha 2, Coluna 2'
    tabela.cell(2, 2).text = 'Linha 2, Coluna 3'